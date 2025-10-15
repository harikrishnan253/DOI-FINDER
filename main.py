"""
Bulk DOI Finder Application
==========================

A FastAPI-based web application that automatically extracts citations from Word documents
and finds their corresponding Digital Object Identifiers (DOIs) using PubMed and CrossRef APIs.

Features:
- Extract citations from .docx documents
- Look up DOIs using PubMed and CrossRef APIs
- Support for APA and AMA citation formats
- Interactive review and editing interface
- Export results to CSV
- Generate documents with embedded DOIs

Author: DOI Finder Team
Version: 1.0
"""

import os
import re
import io
import uuid
import shutil
import logging
import tempfile
import asyncio
from datetime import datetime
from typing import List, Dict, Optional
import certifi
import httpx
from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
import xml.etree.ElementTree as ET
import csv

# =============================================================================
# CONFIGURATION & LOGGING SETUP
# =============================================================================

# Configure logging for the application
logger = logging.getLogger("bulk_doi")
logging.basicConfig(
    level=logging.INFO, 
    format="%(asctime)s %(levelname)s %(message)s"
)

# Application configuration constants
MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB maximum file size
CURRENT_YEAR = datetime.now().year    # Current year for citation validation

# =============================================================================
# FASTAPI APPLICATION SETUP
# =============================================================================

# Initialize FastAPI application
app = FastAPI(
    title="Bulk DOI Finder",
    description="Automatically extract and format academic citations with DOIs",
    version="1.0.0"
)

# Mount static files and templates directories
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# In-memory job storage for processing status and results
# In production, this should be replaced with a proper database
jobs = {}

# =============================================================================
# CITATION EXTRACTION FUNCTIONS
# =============================================================================

def extract_references_section(text: str) -> str:
    """
    Extract the references section from document text.
    
    This function uses multiple regex patterns to identify and extract the
    references section from academic documents. It tries various common
    patterns for references sections and falls back to using the last
    30% of the document if no clear references section is found.
    
    Args:
        text (str): Full document text to search for references
        
    Returns:
        str: Extracted references section text
        
    Note:
        The function looks for patterns like "References", "Bibliography",
        and handles various formatting styles common in academic papers.
    """
    
    # Common patterns for references section headers
    # These patterns account for different formatting styles and variations
    ref_patterns = [
        # Pattern 1: References followed by content until appendix/acknowledgment
        r"(?i)\n\s*references?\s*\n(.*?)(?=\n\s*(?:appendix|acknowledgment|table|figure|author))",
        # Pattern 2: Bibliography followed by content until appendix/acknowledgment  
        r"(?i)\n\s*bibliography\s*\n(.*?)(?=\n\s*(?:appendix|acknowledgment|table|figure|author))",
        # Pattern 3: References at end of document
        r"(?i)\n\s*references?\s*\n(.*?)$",
        # Pattern 4: Bibliography at end of document
        r"(?i)\n\s*bibliography\s*\n(.*?)$",
    ]
    
    # Try each pattern to find the references section
    for pattern in ref_patterns:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            ref_text = match.group(1).strip()
            # Only return if we found a substantial references section
            if len(ref_text) > 100:  # Minimum length threshold
                logger.info(f"Found references section with {len(ref_text)} characters")
                return ref_text
    
    # Fallback strategy: use the last 30% of the document
    # References are typically located at the end of academic papers
    split_point = int(len(text) * 0.7)
    fallback_text = text[split_point:]
    logger.warning("No clear references section found, using last 30% of document")
    return fallback_text

def extract_doi_from_citation(citation_text: str) -> Optional[str]:
    """
    Extract DOI from citation text if it already contains one.
    
    This function searches for existing DOIs in citation text using various
    common DOI patterns and formats. It handles different DOI representations
    including bare DOIs, DOI: prefixed formats, and full URLs.
    
    Args:
        citation_text (str): Citation text to search for DOI
        
    Returns:
        Optional[str]: Extracted DOI if found, None otherwise
        
    Examples:
        >>> extract_doi_from_citation("Smith, J. (2020). Title. Journal. doi:10.1234/example")
        "10.1234/example"
        >>> extract_doi_from_citation("https://doi.org/10.1234/example")
        "10.1234/example"
    """
    
    # DOI patterns in order of preference (most specific first)
    doi_patterns = [
        # Pattern 1: Bare DOI format (10.xxxx/xxxxx)
        r"10\.\d{4,}/[^\s\)\]\n]+",
        # Pattern 2: DOI with "doi:" prefix (case insensitive)
        r"doi:\s*10\.\d{4,}/[^\s\)\]\n]+",
        # Pattern 3: DOI with "DOI:" prefix (case insensitive)
        r"DOI:\s*10\.\d{4,}/[^\s\)\]\n]+", 
        # Pattern 4: Full DOI URL (dx.doi.org or doi.org)
        r"https?://(?:dx\.)?doi\.org/10\.\d{4,}/[^\s\)\]\n]+",
    ]
    
    # Search for DOI using each pattern
    for pattern in doi_patterns:
        match = re.search(pattern, citation_text, re.IGNORECASE)
        if match:
            doi = match.group(0)
            
            # Clean up DOI by removing prefixes and trailing characters
            # Remove "doi:", "DOI:", or URL prefixes
            doi = re.sub(r"^(doi:|DOI:|https?://(?:dx\.)?doi\.org/)\s*", "", doi, flags=re.IGNORECASE)
            # Remove trailing whitespace, parentheses, brackets, or newlines
            doi = re.sub(r"[\s\)\]\n]+$", "", doi)
            
            logger.debug(f"Extracted DOI from citation: {doi.strip()}")
            return doi.strip()
    
    return None

def split_citations(reference_text: str) -> List[str]:
    """
    Split reference text into individual citations.
    
    This function attempts to parse the references section text and split it into
    individual citation entries. It uses multiple strategies:
    1. Numbered citations (1. Author, 2. Author, etc.)
    2. Bracketed citations ([1] Author, [2] Author, etc.)
    3. Author-year format fallback
    
    Args:
        reference_text (str): Full references section text
        
    Returns:
        List[str]: List of individual citation strings
        
    Note:
        The function validates citations by checking for minimum length (30 chars)
        and presence of publication years (1900-2099 range).
    """
    
    citations = []
    
    # Strategy 1: Try numbered citation patterns
    # These are the most common formats in academic papers
    numbered_patterns = [
        # Pattern 1: Numbered citations (1. Author, 2. Author, etc.)
        r"((?:^|\n)\s*\d+\.\s*.*?)(?=\n\s*\d+\.|\Z)",
        # Pattern 2: Bracket numbered citations ([1] Author, [2] Author, etc.)
        r"((?:^|\n)\s*\[\d+\]\s*.*?)(?=\n\s*\[\d+\]|\Z)",
    ]
    
    for pattern in numbered_patterns:
        matches = re.findall(pattern, reference_text, re.MULTILINE | re.DOTALL)
        if len(matches) >= 3:  # Must find at least 3 citations to be valid
            for match in matches:
                # Clean up the citation by removing numbering and extra whitespace
                clean = re.sub(r"^\s*[\d\[\]]+\.?\s*", "", match.strip())
                clean = re.sub(r"\s+", " ", clean).strip()
                
                # Validate citation quality
                if len(clean) >= 30 and re.search(r"\b(19|20)\d{2}\b", clean):
                    citations.append(clean)
            
            if citations:
                logger.info(f"Found {len(citations)} numbered citations")
                return citations
    
    # Strategy 2: Fallback to author-year pattern
    # Look for lines that appear to be individual citations
    lines = reference_text.split("\n")
    potential_citations = []
    
    for line in lines:
        line = line.strip()
        # Check if line looks like a citation (has author pattern and year)
        if (len(line) >= 30 and 
            re.search(r"[A-Z][a-z]+,\s*[A-Z]", line) and  # Author pattern
            re.search(r"\b(19|20)\d{2}\b", line)):        # Year pattern
            potential_citations.append(line)
    
    if potential_citations:
        logger.info(f"Found {len(potential_citations)} author-year citations")
        return potential_citations
    
    # Strategy 3: Final fallback - return non-empty lines
    logger.warning("Using fallback strategy - returning all non-empty lines")
    return [ln.strip() for ln in lines if ln.strip()]

def extract_citation_year(citation: str) -> Optional[str]:
    """
    Extract publication year from citation text.
    
    This function searches for publication years in various common formats
    used in academic citations. It validates that the year is within a
    reasonable range (1900 to current year).
    
    Args:
        citation (str): Citation text to search for year
        
    Returns:
        Optional[str]: Extracted year as string if found, None otherwise
        
    Examples:
        >>> extract_citation_year("Smith, J. (2020). Title. Journal.")
        "2020"
        >>> extract_citation_year("Author, A. 2019; Title.")
        "2019"
    """
    
    # Year patterns in order of preference (most specific first)
    year_patterns = [
        # Pattern 1: Year in parentheses (2023)
        r"\((\d{4})\)",
        # Pattern 2: Year followed by punctuation (2023; or 2023,)
        r"\b(\d{4})[;,.]",
        # Pattern 3: Any 4-digit year in 1900-2099 range
        r"\b(19|20)\d{2}\b",
    ]
    
    for pattern in year_patterns:
        match = re.search(pattern, citation)
        if match:
            # Extract the year value (group 1 if available, otherwise group 0)
            year = match.group(1) if match.lastindex else match.group(0)
            try:
                year_int = int(year)
                # Validate year is within reasonable range
                if 1900 <= year_int <= CURRENT_YEAR:
                    logger.debug(f"Extracted year {year} from citation")
                    return str(year_int)
            except ValueError:
                continue
    
    return None

def parse_citation(citation_text: str, citation_id: int) -> Dict:
    """
    Parse citation text and create structured citation object.
    
    This function takes raw citation text and creates a structured dictionary
    containing all relevant information for processing, including DOI lookup
    status, metadata extraction, and initial confidence scoring.
    
    Args:
        citation_text (str): Raw citation text from document
        citation_id (int): Unique identifier for this citation
        
    Returns:
        Dict: Structured citation object with the following keys:
            - id: Citation identifier
            - original: Original citation text
            - status: Processing status (pending, has_doi, found, not_found)
            - doi: DOI if found
            - confidence: Confidence score (0.0 to 1.0)
            - metadata: Additional metadata (year, source, etc.)
    """
    
    # Initialize citation structure
    citation = {
        "id": citation_id,
        "original": citation_text.strip(),
        "status": "pending",
        "doi": None,
        "confidence": 0.0,
        "metadata": {}
    }
    
    # Check if DOI already exists in the citation
    existing_doi = extract_doi_from_citation(citation_text)
    if existing_doi:
        citation["doi"] = existing_doi
        citation["status"] = "has_doi"
        citation["confidence"] = 1.0  # High confidence for existing DOIs
        citation["metadata"]["existing_doi"] = True
        logger.debug(f"Citation {citation_id} already has DOI: {existing_doi}")
    
    # Extract publication year for search optimization
    year = extract_citation_year(citation_text)
    if year:
        citation["metadata"]["year"] = year
        logger.debug(f"Citation {citation_id} has year: {year}")
    
    return citation

def extract_citations_from_docx(docx_path: str) -> List[Dict]:
    """
    Extract citations from a Word document (.docx file).
    
    This is the main function that orchestrates the citation extraction process.
    It reads the Word document, extracts the references section, splits it into
    individual citations, and parses each citation into a structured format.
    
    Args:
        docx_path (str): Path to the .docx file to process
        
    Returns:
        List[Dict]: List of structured citation objects
        
    Raises:
        HTTPException: If document processing fails
        
    Note:
        This function handles the complete pipeline from document reading
        to structured citation objects ready for DOI lookup.
    """
    
    try:
        logger.info(f"Starting citation extraction from: {docx_path}")
        
        # Load the Word document
        doc = Document(docx_path)
        
        # Extract all text from paragraphs (preserves structure)
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        logger.debug(f"Extracted {len(full_text)} characters from document")
        
        # Step 1: Find and extract the references section
        ref_section = extract_references_section(full_text)
        
        # Step 2: Split references section into individual citations
        citation_texts = split_citations(ref_section)
        
        # Step 3: Parse each citation into structured format
        citations = []
        for i, citation_text in enumerate(citation_texts, 1):
            citation = parse_citation(citation_text, i)
            citations.append(citation)
        
        logger.info(f"Successfully extracted {len(citations)} citations from document")
        return citations
        
    except Exception as e:
        logger.error(f"Error extracting citations from {docx_path}: {e}")
        raise HTTPException(
            status_code=500, 
            detail=f"Citation extraction failed: {str(e)}"
        )

# =============================================================================
# EXTERNAL API CLIENTS
# =============================================================================

async def search_pubmed(query: str) -> Optional[Dict]:
    """
    Search PubMed database for citation metadata using NCBI E-utilities API.
    
    This function performs a two-step search process:
    1. Search for PMIDs (PubMed IDs) matching the query
    2. Fetch detailed metadata for the first result
    
    Args:
        query (str): Search query (typically citation title or key phrases)
        
    Returns:
        Optional[Dict]: Citation metadata if found, None otherwise
            Contains keys: doi, title, authors, journal, year, source
        
    Note:
        Uses NCBI E-utilities API with proper error handling and timeout.
        Respects API rate limits with delays between requests.
    """
    
    # NCBI E-utilities API endpoints
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    search_url = f"{base_url}esearch.fcgi"  # Search for PMIDs
    fetch_url = f"{base_url}efetch.fcgi"    # Fetch metadata for PMIDs
    
    try:
        # Create HTTP client with SSL verification and timeout
        async with httpx.AsyncClient(timeout=30.0, verify=certifi.where()) as client:
            
            # Step 1: Search for PMIDs matching the query
            search_params = {
                "db": "pubmed",           # Search PubMed database
                "term": query[:500],      # Limit query length to avoid URL issues
                "retmode": "json",        # Return JSON format
                "retmax": "5",            # Limit to 5 results for efficiency
                "sort": "relevance"       # Sort by relevance
            }
            
            search_response = await client.get(search_url, params=search_params)
            search_response.raise_for_status()
            search_data = search_response.json()
            
            # Extract PMIDs from search results
            pmids = search_data.get("esearchresult", {}).get("idlist", [])
            if not pmids:
                logger.debug(f"No PMIDs found for query: {query[:100]}...")
                return None
            
            # Step 2: Fetch detailed metadata for the first (most relevant) result
            fetch_params = {
                "db": "pubmed",
                "id": pmids[0],          # Get details for first PMID
                "retmode": "xml"         # XML format for detailed metadata
            }
            
            fetch_response = await client.get(fetch_url, params=fetch_params)
            fetch_response.raise_for_status()
            
            # Parse XML response to extract metadata
            root = ET.fromstring(fetch_response.text)
            
            # Initialize metadata dictionary
            metadata = {"source": "PubMed"}
            
            # Extract DOI from ArticleIdList (highest priority)
            for article_id in root.findall(".//ArticleId"):
                if article_id.get("IdType") == "doi":
                    metadata["doi"] = article_id.text
                    break
            
            # Extract title (truncate if too long)
            title_elem = root.find(".//ArticleTitle")
            if title_elem is not None and title_elem.text:
                metadata["title"] = title_elem.text[:200]
            
            # Extract authors (limit to first 5 for brevity)
            authors = []
            for author in root.findall(".//Author")[:5]:
                last_name = author.find("LastName")
                first_name = author.find("ForeName")
                if last_name is not None and last_name.text:
                    name = last_name.text
                    if first_name is not None and first_name.text:
                        name += f", {first_name.text}"
                    authors.append(name)
            
            if authors:
                metadata["authors"] = "; ".join(authors)
            
            # Extract journal name
            journal_elem = root.find(".//Journal/Title")
            if journal_elem is not None and journal_elem.text:
                metadata["journal"] = journal_elem.text[:100]
            
            # Extract publication year
            year_elem = root.find(".//PubDate/Year")
            if year_elem is not None and year_elem.text:
                metadata["year"] = year_elem.text
            
            # Only return metadata if we found a DOI
            if metadata.get("doi"):
                logger.debug(f"PubMed found DOI: {metadata['doi']}")
                return metadata
            else:
                logger.debug("PubMed found metadata but no DOI")
                return None
            
    except Exception as e:
        logger.warning(f"PubMed search failed for query '{query[:100]}...': {e}")
        return None

async def search_crossref(query: str) -> Optional[Dict]:
    """
    Search CrossRef database for citation metadata using their REST API.
    
    CrossRef is a major DOI registration agency and provides comprehensive
    metadata for scholarly publications. This function searches their database
    and extracts relevant citation information.
    
    Args:
        query (str): Search query (typically citation title or key phrases)
        
    Returns:
        Optional[Dict]: Citation metadata if found, None otherwise
            Contains keys: doi, title, authors, journal, year, source
        
    Note:
        Uses CrossRef REST API with proper headers and rate limiting.
        Includes required mailto parameter for polite API usage.
    """
    
    # CrossRef REST API endpoint
    base_url = "https://api.crossref.org/works"
    
    try:
        # Create HTTP client with SSL verification and timeout
        async with httpx.AsyncClient(timeout=30.0, verify=certifi.where()) as client:
            
            # Prepare search parameters
            params = {
                "query": query[:300],              # Limit query length
                "rows": "5",                       # Limit results for efficiency
                "mailto": "admin@example.com"     # Required for polite API usage
            }
            
            # Required headers for CrossRef API
            headers = {
                "User-Agent": "DOI-Finder/1.0 (mailto:admin@example.com)"
            }
            
            # Make API request
            response = await client.get(base_url, params=params, headers=headers)
            response.raise_for_status()
            data = response.json()
            
            # Extract results from CrossRef response
            items = data.get("message", {}).get("items", [])
            if not items:
                logger.debug(f"No CrossRef results found for query: {query[:100]}...")
                return None
            
            # Process the first (most relevant) result
            item = items[0]
            metadata = {"source": "CrossRef"}
            
            # Extract DOI (primary goal)
            doi = item.get("DOI")
            if doi:
                metadata["doi"] = doi
            
            # Extract title (may be a list, take first)
            titles = item.get("title", [])
            if titles and titles[0]:
                metadata["title"] = titles[0][:200]
            
            # Extract authors (handle CrossRef author format)
            authors = item.get("author", [])
            if authors:
                author_names = []
                for author in authors[:5]:  # Limit to 5 authors for brevity
                    given = author.get("given", "").strip()
                    family = author.get("family", "").strip()
                    if family:
                        # Format as "Last, First" or just "Last"
                        name = f"{family}, {given}" if given else family
                        author_names.append(name)
                
                if author_names:
                    metadata["authors"] = "; ".join(author_names)
            
            # Extract journal/container title
            container_title = item.get("container-title", [])
            if container_title and container_title[0]:
                metadata["journal"] = container_title[0][:100]
            
            # Extract publication year (handle different date formats)
            published = item.get("published-print") or item.get("published-online")
            if published and "date-parts" in published:
                try:
                    # Extract year from date-parts array [year, month, day]
                    year = published["date-parts"][0][0]
                    metadata["year"] = str(year)
                except (IndexError, TypeError):
                    logger.debug("Could not extract year from CrossRef date")
            
            # Only return metadata if we found a DOI
            if metadata.get("doi"):
                logger.debug(f"CrossRef found DOI: {metadata['doi']}")
                return metadata
            else:
                logger.debug("CrossRef found metadata but no DOI")
                return None
            
    except Exception as e:
        logger.warning(f"CrossRef search failed for query '{query[:100]}...': {e}")
        return None

async def lookup_citation_doi(citation: Dict) -> Dict:
    """
    Look up DOI for a citation using multiple external sources.
    
    This function orchestrates the DOI lookup process by:
    1. Extracting searchable queries from the citation text
    2. Searching PubMed database first (higher confidence)
    3. Searching CrossRef database as fallback
    4. Updating citation with results and confidence scores
    
    Args:
        citation (Dict): Citation object to look up DOI for
        
    Returns:
        Dict: Updated citation object with DOI lookup results
        
    Note:
        Includes rate limiting delays between API calls to respect
        service limits and improve success rates.
    """
    
    # Skip lookup if citation already has a DOI
    if citation["status"] == "has_doi":
        logger.debug(f"Citation {citation['id']} already has DOI, skipping lookup")
        return citation
    
    original_text = citation["original"]
    logger.debug(f"Looking up DOI for citation {citation['id']}")
    
    # Build search queries from citation text
    queries = []
    
    # Strategy 1: Extract potential title using various patterns
    title_patterns = [
        # Pattern 1: Text in quotes (often article titles)
        r'"([^"]{10,})"',
        # Pattern 2: Text after year (common in citations)
        r'\d{4}[;,.]?\s*([A-Z][^.]{10,}?)\.',
        # Pattern 3: Text after period (fallback)
        r'\.?\s*([A-Z][^.]{10,}?)\.',
    ]
    
    for pattern in title_patterns:
        match = re.search(pattern, original_text)
        if match:
            potential_title = match.group(1).strip()
            # Validate title length (not too short, not too long)
            if 10 <= len(potential_title) <= 150:
                queries.append(potential_title)
                logger.debug(f"Extracted title query: {potential_title[:50]}...")
                break
    
    # Strategy 2: Add full text as backup query (truncated)
    queries.append(original_text[:200])
    
    # Search PubMed first (higher confidence results)
    for query in queries:
        if not query.strip():
            continue
        
        logger.debug(f"Searching PubMed with query: {query[:50]}...")
        result = await search_pubmed(query)
        if result and result.get("doi"):
            citation.update({
                "status": "found",
                "doi": result["doi"],
                "confidence": 0.9,  # High confidence for PubMed
                "metadata": result,
                "source": "PubMed"
            })
            logger.info(f"Found DOI via PubMed: {result['doi']}")
            return citation
        
        # Rate limiting: small delay between requests
        await asyncio.sleep(0.5)
    
    # Search CrossRef as fallback
    for query in queries:
        if not query.strip():
            continue
        
        logger.debug(f"Searching CrossRef with query: {query[:50]}...")
        result = await search_crossref(query)
        if result and result.get("doi"):
            citation.update({
                "status": "found",
                "doi": result["doi"],
                "confidence": 0.8,  # Good confidence for CrossRef
                "metadata": result,
                "source": "CrossRef"
            })
            logger.info(f"Found DOI via CrossRef: {result['doi']}")
            return citation
        
        # Rate limiting: small delay between requests
        await asyncio.sleep(0.5)
    
    # No DOI found in either source
    citation.update({
        "status": "not_found",
        "confidence": 0.0,
        "message": "No DOI found in PubMed or CrossRef"
    })
    logger.warning(f"No DOI found for citation {citation['id']}")
    
    return citation

# =============================================================================
# CITATION FORMATTING FUNCTIONS
# =============================================================================

def format_citation_apa(metadata: Dict) -> str:
    """
    Format citation metadata in APA (American Psychological Association) style.
    
    APA style follows the format: Author(s). (Year). Title. Journal, Volume(Issue), pages. DOI.
    This function creates a properly formatted citation string from metadata.
    
    Args:
        metadata (Dict): Citation metadata containing authors, year, title, journal, volume, issue, pages, doi
        
    Returns:
        str: Formatted citation in APA style
        
    Example:
        >>> metadata = {"authors": "Smith, J.", "year": "2020", "title": "Example", "journal": "Journal", "volume": "15", "issue": "3", "pages": "123-145", "doi": "10.1234/example"}
        >>> format_citation_apa(metadata)
        "Smith, J. (2020). Example. *Journal*, 15(3), 123-145. https://doi.org/10.1234/example"
    """
    
    parts = []
    
    # Authors - required component (format: Last, F. M., Last, F. M., & Last, F. M.)
    authors = metadata.get("authors", "")
    if authors:
        # Format authors properly for APA style
        author_list = [a.strip() for a in authors.split(";")]
        formatted_authors = []
        
        for i, author in enumerate(author_list):
            if "," in author:
                # Format as "Last, First Middle" -> "Last, F. M."
                last, first = author.split(",", 1)
                first_parts = first.strip().split()
                initials = " ".join([part[0] + "." for part in first_parts if part])
                formatted_authors.append(f"{last.strip()}, {initials}")
            else:
                formatted_authors.append(author)
        
        # Join authors with commas and ampersand before last author
        if len(formatted_authors) == 1:
            parts.append(formatted_authors[0])
        elif len(formatted_authors) == 2:
            parts.append(f"{formatted_authors[0]} & {formatted_authors[1]}")
        else:
            parts.append(", ".join(formatted_authors[:-1]) + f", & {formatted_authors[-1]}")
    
    # Year in parentheses with period - required component
    year = metadata.get("year", "")
    if year:
        parts.append(f"({year}).")
    
    # Title - sentence case (only first word, first word after colon, and proper nouns capitalized)
    title = metadata.get("title", "")
    if title:
        # Convert to proper sentence case for APA
        # Split by colon to handle subtitle capitalization
        if ':' in title:
            main_title, subtitle = title.split(':', 1)
            # Process main title - convert to sentence case
            main_words = main_title.split()
            if main_words:
                main_words[0] = main_words[0].capitalize()
                for i, word in enumerate(main_words[1:], 1):
                    # Convert to lowercase for sentence case (except proper nouns)
                    # For now, convert all words to lowercase for proper sentence case
                    main_words[i] = word.lower()
            
            # Process subtitle - capitalize first word, lowercase rest
            subtitle_words = subtitle.split()
            if subtitle_words:
                subtitle_words[0] = subtitle_words[0].capitalize()
                for i, word in enumerate(subtitle_words[1:], 1):
                    # Convert to lowercase for sentence case (except proper nouns)
                    # For now, convert all words to lowercase for proper sentence case
                    subtitle_words[i] = word.lower()
            
            title_text = " ".join(main_words) + ": " + " ".join(subtitle_words)
        else:
            # No colon, process as single title
            title_words = title.split()
            if title_words:
                title_words[0] = title_words[0].capitalize()
                for i, word in enumerate(title_words[1:], 1):
                    # Convert to lowercase for sentence case (except proper nouns)
                    # For now, convert all words to lowercase for proper sentence case
                    title_words[i] = word.lower()
            title_text = " ".join(title_words)
        
        # Clean up any double periods and ensure only one period at the end
        title_text = title_text.replace("..", ".")
        # Remove any existing period at the end and add exactly one
        title_text = title_text.rstrip('.') + "."
        parts.append(title_text)
    
    # Journal name - italicized in APA style with volume, issue, and pages
    journal = metadata.get("journal", "")
    volume = metadata.get("volume", "")
    issue = metadata.get("issue", "")
    pages = metadata.get("pages", "")
    
    if journal:
        # Journal name italicized (preserve original capitalization)
        journal_part = f"*{journal}*"
        
        # Add volume and issue if available
        if volume:
            if issue:
                journal_part += f", {volume}({issue})"
            else:
                journal_part += f", {volume}"
        
        # Add pages if available
        if pages:
            journal_part += f", {pages}"
        
        parts.append(journal_part + ".")
    
    # DOI as clickable link
    doi = metadata.get("doi", "")
    if doi:
        # Ensure DOI has proper format
        if not doi.startswith("https://doi.org/"):
            if doi.startswith("doi:"):
                doi = doi[4:]  # Remove "doi:" prefix
            parts.append(f"https://doi.org/{doi}")
        else:
            parts.append(doi)
    
    # Join all parts with spaces and add final period, or return error message if incomplete
    if parts:
        citation = " ".join(parts)
        # Ensure citation ends with a period
        if not citation.endswith('.'):
            citation += '.'
        return citation
    else:
        return "Incomplete citation data"

def format_citation_ama(metadata: Dict) -> str:
    """
    Format citation metadata in AMA (American Medical Association) style.
    
    AMA style follows the format: Author(s). Title. Journal. Year;Volume(Issue):pages. doi:DOI.
    This function creates a properly formatted citation string from metadata.
    
    Args:
        metadata (Dict): Citation metadata containing authors, year, title, journal, volume, issue, pages, doi
        
    Returns:
        str: Formatted citation in AMA style
        
    Note:
        AMA style uses abbreviated author names (Last F) and limits to 6 authors
        before adding "et al". Journal names are italicized and abbreviated.
        
    Example:
        >>> metadata = {"authors": "Smith, John", "year": "2020", "title": "Example", "journal": "Journal", "volume": "15", "issue": "3", "pages": "123-145", "doi": "10.1234/example"}
        >>> format_citation_ama(metadata)
        "Smith J. Example. *Journal*. 2020;15(3):123-145. doi:10.1234/example"
    """
    
    parts = []
    
    # Authors in AMA format (Last F, Last F, et al) - up to 6 authors
    authors = metadata.get("authors", "")
    if authors:
        # Split authors and convert to AMA format
        author_list = [a.strip() for a in authors.split(";")]
        ama_authors = []
        
        # Process up to 6 authors (AMA standard)
        for author in author_list[:6]:
            if "," in author:
                # Format as "Last, First Middle" -> "Last F"
                last, first = author.split(",", 1)
                first_parts = first.strip().split()
                first_initial = first_parts[0][:1] if first_parts else ""
                ama_authors.append(f"{last.strip()} {first_initial}")
            else:
                # Keep as-is if no comma found
                ama_authors.append(author)
        
        # Add "et al" if more than 6 authors
        if len(author_list) > 6:
            ama_authors.append("et al")
        
        parts.append(", ".join(ama_authors) + ".")
    
    # Title (not italicized in AMA, sentence case)
    title = metadata.get("title", "")
    if title:
        # Convert to sentence case
        title_words = title.split()
        if title_words:
            title_words[0] = title_words[0].capitalize()
        parts.append(" ".join(title_words) + ".")
    
    # Journal name (italicized and abbreviated)
    journal = metadata.get("journal", "")
    if journal:
        # Basic journal abbreviation (in real implementation, would use a lookup table)
        journal_abbrev = journal
        parts.append(f"*{journal_abbrev}*.")
    
    # Year, volume, issue, and pages in AMA format: Year;Volume(Issue):pages
    year = metadata.get("year", "")
    volume = metadata.get("volume", "")
    issue = metadata.get("issue", "")
    pages = metadata.get("pages", "")
    
    if year:
        citation_info = year
        
        if volume:
            if issue:
                citation_info += f";{volume}({issue})"
            else:
                citation_info += f";{volume}"
            
            if pages:
                citation_info += f":{pages}"
        
        parts.append(citation_info + ".")
    
    # DOI (AMA format: doi:DOI)
    doi = metadata.get("doi", "")
    if doi:
        # Ensure DOI has proper format
        if doi.startswith("https://doi.org/"):
            doi = doi[16:]  # Remove "https://doi.org/" prefix
        elif doi.startswith("doi:"):
            doi = doi[4:]  # Remove "doi:" prefix
        parts.append(f"doi:{doi}")
    
    # Join all parts with spaces, or return error message if incomplete
    return " ".join(parts) if parts else "Incomplete citation data"

# =============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# =============================================================================

def apply_dois_to_document(original_path: str, citations: List[Dict], 
                         apply_mode: str, citation_style: str) -> str:
    """
    Apply DOIs to the original document and create a modified version.
    
    This function takes the original Word document and applies the selected
    citations with their DOIs in the specified format and style. It supports
    multiple application modes for different user preferences.
    
    Args:
        original_path (str): Path to the original .docx file
        citations (List[Dict]): List of citation objects with DOI information
        apply_mode (str): How to apply citations ("append_new_section" or "replace_references")
        citation_style (str): Citation format ("APA" or "AMA")
        
    Returns:
        str: Path to the modified document with DOIs applied
        
    Raises:
        HTTPException: If document processing fails
        ValueError: If no citations are selected for application
        
    Note:
        Creates a new file with "_with_dois" suffix to preserve the original.
        Only processes citations that are marked as "accepted" by the user.
    """
    
    try:
        logger.info(f"Applying DOIs to document: {original_path}")
        
        # Load the original Word document
        doc = Document(original_path)
        
        # Step 1: Format accepted citations with DOIs
        formatted_citations = []
        
        for citation in citations:
            # Only process citations that are accepted and have DOIs
            if citation.get("accepted", False) and citation.get("doi"):
                metadata = citation.get("metadata", {})
                metadata["doi"] = citation["doi"]
                
                # Format citation according to selected style
                if citation_style == "APA":
                    formatted = format_citation_apa(metadata)
                else:  # AMA
                    formatted = format_citation_ama(metadata)
                
                # Add citation number and formatted text
                formatted_citations.append(f"{citation['id']}. {formatted}")
        
        if not formatted_citations:
            raise ValueError("No citations selected for application")
        
        # Step 2: Apply citations based on selected mode
        if apply_mode == "append_new_section":
            # Safe mode: Add new references section at the end
            logger.info("Adding new references section")
            doc.add_heading('References', level=1)
            for formatted_citation in formatted_citations:
                doc.add_paragraph(formatted_citation)
        
        elif apply_mode == "replace_references":
            # Advanced mode: Replace existing references section
            logger.info("Replacing existing references section")
            
            # Find the start of the references section
            ref_start = -1
            for i, paragraph in enumerate(doc.paragraphs):
                if re.search(r"(?i)^\s*(references?|bibliography)\s*$", paragraph.text.strip()):
                    ref_start = i
                    break
            
            if ref_start >= 0:
                # Remove existing references content
                paragraphs_to_remove = []
                for i in range(ref_start + 1, len(doc.paragraphs)):
                    if (doc.paragraphs[i].text.strip() and 
                        not re.search(r"(?i)^\s*(appendix|acknowledgment)", doc.paragraphs[i].text)):
                        paragraphs_to_remove.append(i)
                    else:
                        break
                
                # Remove paragraphs in reverse order to maintain indices
                for i in reversed(paragraphs_to_remove):
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
                
                # Add new formatted citations
                for formatted_citation in formatted_citations:
                    doc.add_paragraph(formatted_citation)
            else:
                # Fallback: append new section if no references found
                logger.warning("No existing references section found, appending new section")
                doc.add_heading('References', level=1)
                for formatted_citation in formatted_citations:
                    doc.add_paragraph(formatted_citation)
        
        # Step 3: Save the modified document
        output_path = original_path.replace(".docx", "_with_dois.docx")
        doc.save(output_path)
        
        logger.info(f"Successfully applied {len(formatted_citations)} DOIs to document")
        return output_path
        
    except Exception as e:
        logger.error(f"Error applying DOIs to document: {e}")
        raise HTTPException(
            status_code=500, 
            detail=f"Document processing failed: {str(e)}"
        )

# =============================================================================
# FASTAPI ROUTES & API ENDPOINTS
# =============================================================================

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """
    Serve the main upload page.
    
    This is the entry point of the application, serving the HTML template
    for file upload and initial configuration.
    
    Args:
        request (Request): FastAPI request object
        
    Returns:
        HTMLResponse: Rendered upload page
    """
    return templates.TemplateResponse("upload.html", {"request": request})

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), citation_format: str = Form("APA")):
    """
    Handle file upload and initiate document processing.
    
    This endpoint receives uploaded Word documents, validates them,
    saves them to temporary storage, and starts the background
    processing task for citation extraction and DOI lookup.
    
    Args:
        file (UploadFile): The uploaded .docx file
        citation_format (str): Preferred citation format (APA or AMA)
        
    Returns:
        Dict: Job information including job_id and status
        
    Raises:
        HTTPException: If file validation fails or upload errors occur
        
    Note:
        Processing happens asynchronously in the background. The client
        should poll the job status endpoint to check progress.
    """
    
    # Validate uploaded file
    if not file.filename.endswith('.docx'):
        raise HTTPException(
            status_code=400, 
            detail="Only .docx files are supported"
        )
    
    if file.size and file.size > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=400, 
            detail=f"File too large. Max size: {MAX_UPLOAD_BYTES // (1024*1024)}MB"
        )
    
    # Generate unique job identifier
    job_id = str(uuid.uuid4())
    
    # Ensure temp directory exists and save uploaded file
    os.makedirs("temp", exist_ok=True)
    file_path = f"temp/{job_id}_{file.filename}"
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"File upload failed: {str(e)}"
        )
    
    # Create job record for tracking
    jobs[job_id] = {
        "id": job_id,
        "filename": file.filename,
        "filepath": file_path,
        "status": "uploaded",
        "citation_format": citation_format,
        "created_at": datetime.now().isoformat(),
        "citations": []
    }
    
    logger.info(f"File uploaded successfully: {file.filename} (Job: {job_id})")
    
    # Start background processing task
    asyncio.create_task(process_document(job_id))
    
    return {
        "job_id": job_id, 
        "status": "uploaded", 
        "message": "Extended processing started - this may take up to 10 minutes for complete results"
    }


@app.get("/process/{job_id}")
async def process_status(job_id: str):
    """
    Get the current processing status for a job.
    
    This endpoint allows clients to check the progress of document processing,
    including citation extraction and DOI lookup progress.
    
    Args:
        job_id (str): Unique job identifier
        
    Returns:
        Dict: Status information including current status and progress percentage
        
    Raises:
        HTTPException: If job ID is not found
    """
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    return {"status": job["status"], "progress": job.get("progress", 0)}

async def process_document(job_id: str):
    """
    Extended background task to process ALL citations completely.
    
    This is the main processing function that runs asynchronously after file upload.
    It orchestrates the complete pipeline: citation extraction, DOI lookup,
    and result compilation. Designed to handle large documents with many citations.
    
    Args:
        job_id (str): Unique job identifier for tracking progress
        
    Note:
        This function updates the job status and progress in real-time,
        allowing the frontend to display progress updates to users.
        Includes extended timeouts to handle large documents completely.
    """
    
    try:
        job = jobs[job_id]
        job["status"] = "processing"
        job["progress"] = 10
        
        # Extract citations
        logger.info(f"Starting citation extraction for job {job_id}")
        citations = extract_citations_from_docx(job["filepath"])
        job["citations"] = citations
        job["progress"] = 30
        
        total_citations = len(citations)
        logger.info(f"Found {total_citations} citations to process for job {job_id}")
        
        # Look up DOIs for each citation with EXTENDED processing
        processed = 0
        successful_lookups = 0
        
        # EXTENDED timeout for complete processing - 10 minutes total
        start_time = datetime.now()
        timeout_minutes = 10
        
        for i, citation in enumerate(citations):
            # Check timeout but allow more generous time
            elapsed_minutes = (datetime.now() - start_time).seconds / 60
            if elapsed_minutes > timeout_minutes:
                logger.warning(f"Extended timeout reached for job {job_id} after {elapsed_minutes:.1f} minutes")
                # Mark remaining citations as not found
                for remaining_citation in citations[processed:]:
                    remaining_citation["status"] = "not_found"
                    remaining_citation["confidence"] = 0.0
                    remaining_citation["message"] = "Processing timeout"
                break
            
            try:
                logger.info(f"Looking up DOI for citation {i+1}/{total_citations} in job {job_id}")
                
                # Extended timeout per citation - 45 seconds
                await asyncio.wait_for(
                    lookup_citation_doi(citation), 
                    timeout=45.0
                )
                
                if citation.get("doi"):
                    successful_lookups += 1
                    logger.info(f"Successfully found DOI for citation {i+1}: {citation['doi']}")
                
            except asyncio.TimeoutError:
                citation["status"] = "not_found"
                citation["confidence"] = 0.0
                citation["message"] = "DOI lookup timeout (45s)"
                logger.warning(f"Timeout on citation {i+1} for job {job_id}")
                
            except Exception as e:
                citation["status"] = "not_found"
                citation["confidence"] = 0.0
                citation["message"] = f"Lookup error: {str(e)}"
                logger.error(f"Error on citation {i+1} for job {job_id}: {e}")
            
            processed += 1
            progress_percent = 30 + ((processed / total_citations) * 65)
            job["progress"] = min(95, progress_percent)
            
            # Log detailed progress every 5 citations
            if processed % 5 == 0 or processed == total_citations:
                logger.info(f"Job {job_id} progress: {processed}/{total_citations} citations processed, {successful_lookups} DOIs found")
        
        job["status"] = "completed"
        job["progress"] = 100
        job["completed_at"] = datetime.now().isoformat()
        
        final_stats = {
            "total": total_citations,
            "processed": processed,
            "dois_found": successful_lookups,
            "processing_time": (datetime.now() - start_time).seconds
        }
        
        job["final_stats"] = final_stats
        
        logger.info(f"Job {job_id} completed successfully:")
        logger.info(f"  - Total citations: {total_citations}")
        logger.info(f"  - Citations processed: {processed}")
        logger.info(f"  - DOIs found: {successful_lookups}")
        logger.info(f"  - Processing time: {final_stats['processing_time']} seconds")
        
    except Exception as e:
        logger.error(f"Critical error processing job {job_id}: {e}")
        job["status"] = "error"
        job["error"] = str(e)
        job["progress"] = 0

@app.get("/job/{job_id}")
async def get_job_status(job_id: str):
    """
    Get comprehensive job status and results.
    
    This endpoint provides detailed information about a job's progress,
    including citation statistics and processing results.
    
    Args:
        job_id (str): Unique job identifier
        
    Returns:
        Dict: Complete job information including citations and statistics
        
    Raises:
        HTTPException: If job ID is not found
    """
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    
    # Calculate citation statistics for dashboard display
    citations = job.get("citations", [])
    stats = {
        "total": len(citations),
        "has_doi": len([c for c in citations if c["status"] == "has_doi"]),
        "found": len([c for c in citations if c["status"] == "found"]),
        "not_found": len([c for c in citations if c["status"] == "not_found"]),
        "pending": len([c for c in citations if c["status"] == "pending"])
    }
    
    return {
        "job": job,
        "stats": stats
    }

@app.get("/review/{job_id}", response_class=HTMLResponse)
async def review_results(request: Request, job_id: str):
    """
    Serve the review page for job results.
    
    This endpoint renders the interactive review page where users can
    examine extracted citations, edit DOIs, select citations for
    application, and download processed documents.
    
    Args:
        request (Request): FastAPI request object
        job_id (str): Unique job identifier
        
    Returns:
        HTMLResponse: Rendered review page with job data
        
    Raises:
        HTTPException: If job ID is not found
    """
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    citations = job.get("citations", [])
    
    # Calculate statistics for the review page
    stats = {
        "total": len(citations),
        "has_doi": len([c for c in citations if c["status"] == "has_doi"]),
        "found": len([c for c in citations if c["status"] == "found"]),
        "not_found": len([c for c in citations if c["status"] == "not_found"])
    }
    
    return templates.TemplateResponse("review.html", {
        "request": request,
        "job": job,
        "citations": citations,
        "stats": stats
    })

@app.post("/apply/{job_id}")
async def apply_dois(job_id: str, request: Request):
    """
    Apply selected DOIs to the document and generate output file.
    
    This endpoint processes user selections and edits, applies the selected
    citations with DOIs to the original document, and prepares it for download.
    
    Args:
        job_id (str): Unique job identifier
        request (Request): FastAPI request containing user selections and edits
        
    Returns:
        Dict: Success status and download URL
        
    Raises:
        HTTPException: If job not found or processing fails
    """
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Parse request data
    data = await request.json()
    apply_mode = data.get("apply_mode", "append_new_section")
    citation_style = data.get("citation_style", "APA")
    selected_citations = data.get("selected_citations", [])
    citation_updates = data.get("citation_updates", {})
    
    job = jobs[job_id]
    citations = job["citations"]
    
    # Apply user edits to citations
    for citation in citations:
        cid = str(citation["id"])
        citation["accepted"] = cid in selected_citations
        
        # Handle manual DOI edits
        if cid in citation_updates and citation_updates[cid]:
            citation["doi"] = citation_updates[cid]
            if citation["status"] == "not_found":
                citation["status"] = "found"
                citation["confidence"] = 0.5  # User-provided DOI
    
    try:
        # Generate document with applied DOIs
        output_path = apply_dois_to_document(
            job["filepath"],
            citations,
            apply_mode,
            citation_style
        )
        
        job["output_path"] = output_path
        return {"status": "success", "download_url": f"/download/{job_id}"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{job_id}")
async def download_result(job_id: str):
    """
    Download the processed document with applied DOIs.
    
    This endpoint serves the final Word document with all selected
    citations and their DOIs applied in the user's chosen format.
    
    Args:
        job_id (str): Unique job identifier
        
    Returns:
        FileResponse: Processed Word document for download
        
    Raises:
        HTTPException: If job not found or file not available
    """
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    if "output_path" not in job or not os.path.exists(job["output_path"]):
        raise HTTPException(status_code=404, detail="Processed file not found")
    
    return FileResponse(
        path=job["output_path"],
        filename=job["filename"].replace(".docx", "_with_dois.docx"),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

from fastapi.responses import Response

@app.get("/export/{job_id}")
async def export_csv(job_id: str):
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")

    job = jobs[job_id]
    citations = job.get("citations", [])

    csv_data = io.StringIO()
    writer = csv.writer(csv_data)

    # Write CSV header
    writer.writerow([
        "ID", "Original Citation", "Status", "DOI", "Confidence",
        "Title", "Authors", "Journal", "Year", "Source"
    ])

    # Write citation rows
    for citation in citations:
        metadata = citation.get("metadata", {})
        writer.writerow([
            citation["id"],
            citation["original"],
            citation["status"],
            citation.get("doi", ""),
            citation.get("confidence", ""),
            metadata.get("title", ""),
            metadata.get("authors", ""),
            metadata.get("journal", ""),
            metadata.get("year", ""),
            citation.get("source", "")
        ])

    csv_content = csv_data.getvalue()
    csv_data.close()

    return Response(
        content=csv_content,
        media_type="text/csv",
        headers={
            "Content-Disposition": f"attachment; filename=citations_{job_id}.csv"
        }
    )


@app.get("/health")
async def health_check():
    """
    Health check endpoint for monitoring and load balancers.
    
    Returns:
        Dict: Application health status and timestamp
    """
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# =============================================================================
# APPLICATION STARTUP
# =============================================================================

if __name__ == "__main__":
    """
    Start the FastAPI application server.
    
    This runs the application in development mode with hot reloading.
    For production deployment, use a proper ASGI server like Gunicorn.
    """
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)